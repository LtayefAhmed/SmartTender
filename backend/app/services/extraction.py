"""Document text extraction, with OCR fallback.

This is what turns a stored PDF into something the scoring engine can actually
read. Without it, ``field_of_work`` and ``keywords`` — 55% of the total weight
— evaluate a tender on its title alone, and the CV-matching module downstream
has no requirements to match against.

The strategy is cheapest-first, the same principle as duplicate detection:

    1. **Digital text layer** (pypdf / python-docx / BeautifulSoup).
       Pure Python, milliseconds, no system dependency. Handles the large
       majority of tender documents, which are produced digitally.

    2. **OCR** (pypdfium2 → OpenCV → Tesseract), *per page*, only where the
       text layer came back essentially empty.

Doing the fallback per page rather than per document matters: tender files are
routinely hybrid — a digital cover page followed by a scanned, stamped annex.
Choosing one strategy for the whole file either wastes a second per page on
pages that did not need it, or silently loses the pages that did.

Every failure degrades rather than raises. A document that cannot be read
leaves the tender with less text, never without a tender.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import Any

from app.core.config import get_settings
from app.core.logging import get_logger
from app.core.metrics import observe_stage

logger = get_logger(__name__)

__all__ = [
    "DOCUMENT_MARKER",
    "DocumentExtractor",
    "ExtractedText",
    "document_priority",
    "get_extractor",
]

#: Filename fragments that mark the substantive pieces of a French public
#: procurement dossier. CCTP, CCAP and the règlement de consultation carry the
#: requirements; ATTRI, DC1 and DC2 are forms a bidder fills in rather than
#: reads. Names are the only signal available before a file is opened, and
#: French procurement is regular enough about them to make this worth using.
_IMPORTANT_MARKERS = (
    "cctp", "ccap", "cdc", "cahier", "reglement", "règlement", "_rc", "rc_",
    "dce", "consultation", "technique", "cahier des charges", "annexe technique",
)


#: Opens each document inside a merged block of text. Two consumers rely on it:
#: a tender merges its attachments, and an archive merges its members. Chosen
#: to be unmistakable in prose and trivial to split on — downstream, a
#: 440 000-character tender is far too long to embed as one vector, so it must
#: be chunked, and a chunk is only useful if you know whether it came from the
#: CCTP or from a privacy notice nobody reads.
DOCUMENT_MARKER = "\n===== DOCUMENT:"
_DOCUMENT_MARKER = DOCUMENT_MARKER


@dataclass(slots=True)
class _OcrBudget:
    """Pages of OCR still allowed, shared across one extraction.

    ``max_ocr_pages`` bounds a single document, which was enough while a
    document was the unit of work. An archive changed that: a dossier holding
    twenty scanned PDFs would spend twenty times the intended budget, and one
    consultation could occupy an OCR worker for the better part of an hour.
    The bound has to be shared by everything one call touches.
    """

    remaining: int

    def take(self, wanted: int) -> int:
        granted = max(0, min(wanted, self.remaining))
        self.remaining -= granted
        return granted


def document_priority(name: str | None) -> int:
    """0 for a substantive document, 1 for anything else.

    Used wherever a cap can bind — the attachment limit, the archive member
    limit — so that what gets dropped is the administrative filler and not the
    specification.
    """
    lowered = (name or "").lower()
    return 0 if any(marker in lowered for marker in _IMPORTANT_MARKERS) else 1

_WHITESPACE = re.compile(r"[ \t ]+")
_BLANK_LINES = re.compile(r"\n{3,}")
#: Ligatures and dashes that PDF producers emit and that break keyword matching.
_PDF_ARTEFACTS = {
    "ﬀ": "ff", "ﬁ": "fi", "ﬂ": "fl", "ﬃ": "ffi", "ﬄ": "ffl",
    "‘": "'", "’": "'", "“": '"', "”": '"',
    "–": "-", "—": "-", "­": "",
}


@dataclass(slots=True)
class ExtractedText:
    """What one document yielded."""

    text: str = ""
    method: str = "none"          # digital | ocr | mixed | none
    pages_total: int = 0
    pages_ocr: int = 0
    truncated: bool = False
    error: str | None = None
    warnings: list[str] = field(default_factory=list)

    @property
    def ok(self) -> bool:
        return bool(self.text.strip())

    @property
    def char_count(self) -> int:
        return len(self.text)

    def to_dict(self) -> dict[str, Any]:
        return {
            "method": self.method,
            "chars": self.char_count,
            "pages_total": self.pages_total,
            "pages_ocr": self.pages_ocr,
            "truncated": self.truncated,
            "error": self.error,
            "warnings": self.warnings,
        }


def clean_extracted_text(raw: str, *, max_chars: int | None = None) -> tuple[str, bool]:
    """Normalise extracted text for storage and scoring."""
    if not raw:
        return "", False

    text = raw
    for bad, good in _PDF_ARTEFACTS.items():
        text = text.replace(bad, good)

    # PDF extraction hyphenates across line breaks; rejoining is what makes
    # "développe-\nment" match the keyword "développement".
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _WHITESPACE.sub(" ", text)
    text = "\n".join(line.strip() for line in text.split("\n"))
    text = _BLANK_LINES.sub("\n\n", text).strip()

    truncated = False
    if max_chars and len(text) > max_chars:
        text = text[:max_chars].rsplit(" ", 1)[0]
        truncated = True
    return text, truncated


class DocumentExtractor:
    """Extracts readable text from PDF, DOCX and HTML documents."""

    def __init__(self) -> None:
        settings = get_settings().extraction
        self.enabled = settings.enabled
        self.ocr_enabled = settings.ocr_enabled
        self.tesseract_cmd = settings.tesseract_cmd
        self.ocr_languages = settings.ocr_languages
        self.min_chars_before_ocr = settings.min_chars_before_ocr
        self.max_ocr_pages = settings.max_ocr_pages
        self.ocr_dpi = settings.ocr_dpi
        self.max_pdf_pages = settings.max_pdf_pages
        self.max_chars = settings.max_chars_per_document
        self.max_document_bytes = settings.max_document_bytes
        self.archive_max_members = settings.archive_max_members
        self.archive_max_total_bytes = settings.archive_max_total_bytes
        self.archive_max_depth = settings.archive_max_depth
        self._ocr_available: bool | None = None

    # ------------------------------------------------------------------
    def extract(self, content: bytes, *, content_type: str | None = None,
                filename: str | None = None, _depth: int = 0,
                _ocr_budget: _OcrBudget | None = None) -> ExtractedText:
        """Extract text from a document, dispatching on its real type."""
        if not self.enabled:
            return ExtractedText(method="none", error="Extraction is disabled.")
        if not content:
            return ExtractedText(method="none", error="Document is empty.")
        if len(content) > self.max_document_bytes:
            return ExtractedText(
                method="none",
                error=f"Document exceeds the {self.max_document_bytes} byte extraction limit.",
            )

        # A top-level document gets the full per-document allowance; members of
        # an archive share one, handed down by the caller.
        budget = _ocr_budget or _OcrBudget(self.max_ocr_pages)

        kind = self._detect(content, content_type, filename)
        with observe_stage("extract"):
            if kind == "pdf":
                return self._extract_pdf(content, budget)
            if kind == "docx":
                return self._extract_docx(content)
            if kind == "html":
                return self._extract_html(content)
            if kind == "text":
                text, truncated = clean_extracted_text(
                    content.decode("utf-8", errors="replace"), max_chars=self.max_chars
                )
                return ExtractedText(text=text, method="digital", truncated=truncated)
            if kind == "zip":
                return self._extract_archive(content, depth=_depth, budget=budget)
            if kind == "7z":
                return self._extract_7z(content, depth=_depth, budget=budget)
            if kind == "xlsx":
                return self._extract_xlsx(content)

        return ExtractedText(method="none", error=f"Unsupported document type '{kind}'.")

    @staticmethod
    def _detect(content: bytes, content_type: str | None, filename: str | None) -> str:
        """Determine the real type. Content wins over the declared type."""
        from app.services.validation import sniff_mime

        mime = sniff_mime(content[:8192], full=content)
        if mime == "application/pdf":
            return "pdf"
        if "wordprocessingml" in mime:
            return "docx"
        if mime in {"text/html"}:
            return "html"
        if mime == "text/plain":
            # A declared .html with a text/plain sniff is still HTML.
            suffix = (filename or "").lower()
            return "html" if suffix.endswith((".html", ".htm")) else "text"
        # Checked after the OOXML formats above, which are themselves ZIPs: a
        # .docx reaching the archive branch would be read as a bag of XML parts
        # instead of a document.
        if mime in {"application/zip", "application/x-zip-compressed"}:
            return "zip"
        if "spreadsheetml" in mime:
            return "xlsx"
        # 7-Zip has no registered MIME of its own in most sniffers, so it
        # arrives as octet-stream; the six-byte signature is unambiguous.
        if content[:6] == b"7z\xbc\xaf\x27\x1c":
            return "7z"
        return mime

    # ------------------------------------------------------------------
    def _extract_pdf(self, content: bytes, budget: _OcrBudget) -> ExtractedText:
        """Digital text first, per-page OCR only where it came back empty."""
        try:
            from pypdf import PdfReader
        except ImportError:  # pragma: no cover - declared dependency
            return ExtractedText(method="none", error="pypdf is not installed.")

        warnings: list[str] = []
        try:
            reader = PdfReader(io.BytesIO(content))
            if reader.is_encrypted:
                # Many tender PDFs carry an owner password with an empty user
                # password; that decrypts silently and is worth attempting.
                try:
                    reader.decrypt("")
                except Exception:
                    return ExtractedText(
                        method="none", error="PDF is encrypted and could not be opened."
                    )
            pages = reader.pages[: self.max_pdf_pages]
            total_pages = len(reader.pages)
        except Exception as exc:
            return ExtractedText(method="none", error=f"Unreadable PDF: {exc}")

        if total_pages > self.max_pdf_pages:
            warnings.append(
                f"Only the first {self.max_pdf_pages} of {total_pages} pages were read."
            )

        page_texts: list[str] = []
        needs_ocr: list[int] = []
        for index, page in enumerate(pages):
            try:
                text = page.extract_text() or ""
            except Exception:
                text = ""
            page_texts.append(text)
            if len(text.strip()) < self.min_chars_before_ocr:
                needs_ocr.append(index)

        pages_ocr = 0
        if needs_ocr and self.ocr_enabled and self._ensure_ocr():
            granted = budget.take(len(needs_ocr))
            selected = needs_ocr[:granted]
            if granted < len(needs_ocr):
                warnings.append(
                    f"{len(needs_ocr)} pages needed OCR; only {granted} were processed."
                )
            recovered = self._ocr_pages(content, selected) if selected else {}
            for index, text in recovered.items():
                if text.strip():
                    page_texts[index] = text
                    pages_ocr += 1
        elif needs_ocr and self.ocr_enabled:
            warnings.append(
                f"{len(needs_ocr)} page(s) appear scanned but OCR is unavailable."
            )

        joined = "\n\n".join(t for t in page_texts if t and t.strip())
        text, truncated = clean_extracted_text(joined, max_chars=self.max_chars)

        if pages_ocr and pages_ocr < len(page_texts):
            method = "mixed"
        elif pages_ocr:
            method = "ocr"
        elif text:
            method = "digital"
        else:
            method = "none"

        return ExtractedText(
            text=text,
            method=method,
            pages_total=len(pages),
            pages_ocr=pages_ocr,
            truncated=truncated,
            warnings=warnings,
            error=None if text else "No readable text found in the PDF.",
        )

    # ------------------------------------------------------------------
    def _ensure_ocr(self) -> bool:
        """Check once whether Tesseract is genuinely usable."""
        if self._ocr_available is not None:
            return self._ocr_available
        try:
            import pypdfium2  # noqa: F401
            import pytesseract

            if self.tesseract_cmd:
                pytesseract.pytesseract.tesseract_cmd = self.tesseract_cmd
            # Invoke it: an installed wrapper with a missing binary is the
            # common failure, and it only surfaces on a real call.
            pytesseract.get_tesseract_version()
            self._ocr_available = True
            logger.info("extraction.ocr_available", languages=self.ocr_languages)
        except Exception as exc:
            self._ocr_available = False
            logger.warning(
                "extraction.ocr_unavailable",
                error=str(exc),
                hint="Install Tesseract and set SMARTTENDER_EXTRACTION__TESSERACT_CMD.",
            )
        return self._ocr_available

    def _ocr_pages(self, content: bytes, page_indices: list[int]) -> dict[int, str]:
        """Rasterise and OCR the given pages.

        pypdfium2 renders without needing Poppler, which keeps this working
        identically on Windows and in the container.
        """
        results: dict[int, str] = {}
        try:
            import pypdfium2 as pdfium
            import pytesseract
        except ImportError:  # pragma: no cover
            return results

        scale = self.ocr_dpi / 72.0
        document = None
        try:
            document = pdfium.PdfDocument(content)
            for index in page_indices:
                try:
                    page = document[index]
                    bitmap = page.render(scale=scale, grayscale=True)
                    image = bitmap.to_pil()
                    image = self._preprocess(image)
                    results[index] = pytesseract.image_to_string(
                        image, lang=self.ocr_languages
                    )
                except Exception as exc:
                    logger.debug("extraction.ocr_page_failed", page=index, error=str(exc))
        except Exception as exc:
            logger.warning("extraction.ocr_failed", error=str(exc))
        finally:
            if document is not None:
                try:
                    document.close()
                except Exception:
                    pass
        return results

    def _preprocess(self, image: Any) -> Any:
        """Clean a rasterised page before OCR.

        Adaptive thresholding plus a median blur is the standard recipe for
        scanned administrative documents: it removes the grey cast and the
        speckle that scanners and stamps introduce, which is worth several
        percent of character accuracy on exactly the paperwork tenders attach.
        Falls back to the untouched image if OpenCV is unavailable.
        """
        try:
            import cv2
            import numpy as np
            from PIL import Image

            array = np.array(image.convert("L"))
            array = cv2.medianBlur(array, 3)
            array = cv2.adaptiveThreshold(
                array, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 10
            )
            return Image.fromarray(array)
        except Exception:
            return image

    # ------------------------------------------------------------------
    def _extract_docx(self, content: bytes) -> ExtractedText:
        try:
            import docx
        except ImportError:  # pragma: no cover
            return ExtractedText(method="none", error="python-docx is not installed.")

        try:
            document = docx.Document(io.BytesIO(content))
        except Exception as exc:
            return ExtractedText(method="none", error=f"Unreadable DOCX: {exc}")

        parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
        # Requirements in tender documents live in tables far more often than
        # in paragraphs, so skipping them would lose the most valuable content.
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        text, truncated = clean_extracted_text("\n".join(parts), max_chars=self.max_chars)
        return ExtractedText(
            text=text,
            method="digital" if text else "none",
            truncated=truncated,
            error=None if text else "No readable text found in the DOCX.",
        )

    def _extract_html(self, content: bytes) -> ExtractedText:
        try:
            from app.connectors.parsing.selectors import parse_html

            soup = parse_html(content)
            for tag in soup(["script", "style", "nav", "footer", "header", "noscript"]):
                tag.decompose()
            raw = soup.get_text("\n", strip=True)
        except Exception as exc:
            return ExtractedText(method="none", error=f"Unreadable HTML: {exc}")

        text, truncated = clean_extracted_text(raw, max_chars=self.max_chars)
        return ExtractedText(
            text=text,
            method="digital" if text else "none",
            truncated=truncated,
            error=None if text else "No readable text found in the HTML.",
        )

    def _extract_archive(self, content: bytes, *, depth: int,
                         budget: _OcrBudget) -> ExtractedText:
        """Read every member of a ZIP, recursively.

        Buyers frequently publish the entire dossier as a single archive, so an
        unopened ZIP can hide the CCTP itself. Recursion is bounded on three
        independent axes because an archive is attacker-controlled input: a few
        kilobytes can expand to gigabytes, an archive can contain itself, and
        the member count is unbounded. Exceeding a bound is reported as a
        warning rather than an error — partial content plus a visible note is
        worth more than a silent nothing.
        """
        import io
        import zipfile

        if depth >= self.archive_max_depth:
            return ExtractedText(
                method="none",
                error=f"Nested archives beyond depth {self.archive_max_depth} are not opened.",
            )

        try:
            archive = zipfile.ZipFile(io.BytesIO(content))
        except Exception as exc:
            return ExtractedText(method="none", error=f"Unreadable archive: {exc}")

        members = [entry for entry in archive.infolist() if not entry.is_dir() and entry.file_size]
        # Most important first, so a binding member cap drops the filler.
        members.sort(key=lambda entry: (document_priority(entry.filename), entry.filename))

        read: list[tuple[str, bytes]] = []
        warnings: list[str] = []
        # Named apart from the OCR budget: they bound different resources and
        # sharing the name once cost an afternoon.
        size_budget = self.archive_max_total_bytes

        if len(members) > self.archive_max_members:
            warnings.append(
                f"Archive holds {len(members)} files; only the "
                f"{self.archive_max_members} most relevant were read."
            )
            members = members[: self.archive_max_members]

        for entry in members:
            if entry.flag_bits & 0x1:
                warnings.append(f"{entry.filename}: encrypted, not read.")
                continue
            # The budget is the real defence against a zip bomb. It is checked
            # against the *declared* size before decompressing, which is safe
            # here only because zipfile stops a read at the declared size: a
            # member cannot expand past its own header, so the running total
            # cannot be exceeded by lying about it.
            if entry.file_size > size_budget:
                warnings.append(f"{entry.filename}: skipped, archive size budget exhausted.")
                continue
            try:
                with archive.open(entry) as handle:
                    payload = handle.read()
            except Exception as exc:
                warnings.append(f"{entry.filename}: unreadable ({type(exc).__name__}).")
                continue

            size_budget -= len(payload)
            read.append((entry.filename, payload))

        return self._merge_members(read, depth=depth, warnings=warnings, budget=budget)

    def _extract_7z(self, content: bytes, *, depth: int,
                    budget: _OcrBudget) -> ExtractedText:
        """Read a 7-Zip dossier.

        Which archive format a buyer used is not a property of the tender, so
        supporting one and not the other would lose documents for an arbitrary
        reason. py7zr has no incremental per-member API, hence the size check
        happens before extraction rather than during it.
        """
        import io

        if depth >= self.archive_max_depth:
            return ExtractedText(
                method="none",
                error=f"Nested archives beyond depth {self.archive_max_depth} are not opened.",
            )
        try:
            import py7zr
        except ImportError:  # pragma: no cover - the container always has it
            return ExtractedText(method="none", error="7z support is not installed.")

        import pathlib
        import tempfile

        warnings: list[str] = []
        read: list[tuple[str, bytes]] = []

        # py7zr dropped its in-memory reader, so members land on disk. The
        # directory is temporary and removed on the way out; nothing survives
        # the call.
        with tempfile.TemporaryDirectory(prefix="smarttender-7z-") as workspace:
            root = pathlib.Path(workspace).resolve()
            try:
                with py7zr.SevenZipFile(io.BytesIO(content)) as archive:
                    if archive.needs_password():
                        return ExtractedText(method="none", error="Archive is encrypted.")
                    entries = [entry for entry in archive.list() if not entry.is_directory]
                    total = sum(entry.uncompressed for entry in entries)
                    if total > self.archive_max_total_bytes:
                        return ExtractedText(
                            method="none",
                            error=(
                                f"Archive expands to {total} bytes, "
                                "above the extraction budget."
                            ),
                        )
                    names = sorted(
                        (entry.filename for entry in entries),
                        key=lambda name: (document_priority(name), name),
                    )
                    if len(names) > self.archive_max_members:
                        warnings.append(
                            f"Archive holds {len(names)} files; only the "
                            f"{self.archive_max_members} most relevant were read."
                        )
                        names = names[: self.archive_max_members]
                    archive.extract(path=workspace, targets=names)
            except Exception as exc:
                return ExtractedText(method="none", error=f"Unreadable archive: {exc}")

            for name in names:
                member = (root / name).resolve()
                # An archive can name a member "../../etc/passwd". Writing
                # outside the workspace is the extraction step's problem;
                # *reading* what landed outside it would be ours.
                if not member.is_relative_to(root) or not member.is_file():
                    warnings.append(f"{name}: skipped, resolved outside the extraction directory.")
                    continue
                try:
                    read.append((name, member.read_bytes()))
                except Exception as exc:
                    warnings.append(f"{name}: unreadable ({type(exc).__name__}).")

        return self._merge_members(read, depth=depth, warnings=warnings, budget=budget)

    def _merge_members(
        self, members: list[tuple[str, bytes]], *, depth: int,
        warnings: list[str], budget: _OcrBudget,
    ) -> ExtractedText:
        """Extract each archive member and merge the results.

        Shared by every archive format so a dossier reads identically whichever
        tool the buyer packed it with.
        """
        collected: list[str] = []
        methods: set[str] = set()
        pages_ocr = 0

        for name, payload in members:
            inner = self.extract(payload, filename=name, _depth=depth + 1, _ocr_budget=budget)
            if inner.ok:
                collected.append(f"{_DOCUMENT_MARKER} {name}\n{inner.text}")
                methods.add(inner.method)
                pages_ocr += inner.pages_ocr
            elif inner.error:
                warnings.append(f"{name}: {inner.error}")
            warnings.extend(f"{name}: {warning}" for warning in inner.warnings)

        text, truncated = clean_extracted_text("\n\n".join(collected), max_chars=self.max_chars)
        method = methods.pop() if len(methods) == 1 else ("mixed" if methods else "none")
        return ExtractedText(
            text=text,
            method=method,
            pages_ocr=pages_ocr,
            truncated=truncated,
            warnings=warnings,
            error=None if text else "No readable text found in the archive.",
        )

    def _extract_xlsx(self, content: bytes) -> ExtractedText:
        """Flatten a spreadsheet to text, sheet by sheet.

        A procurement dossier's spreadsheets are not incidental: the BPU and
        the DQE hold the price structure, and annexes named "Exigences,
        pénalités, livrables et indicateurs" are the requirements matrix
        itself. Formulas are skipped in favour of cached values — a matcher
        needs "99,5 % de disponibilité", not "=B4*C4".
        """
        import io

        try:
            from openpyxl import load_workbook

            workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        except Exception as exc:
            return ExtractedText(method="none", error=f"Unreadable spreadsheet: {exc}")

        lines: list[str] = []
        try:
            for sheet in workbook.worksheets:
                lines.append(f"[{sheet.title}]")
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(cell).strip() for cell in row if cell not in (None, "")]
                    if cells:
                        # Tab-separated: a row is one line, so a requirement and
                        # its target value stay adjacent after chunking.
                        lines.append("\t".join(cells))
        except Exception as exc:
            lines.append(f"[lecture interrompue: {exc}]")
        finally:
            workbook.close()

        text, truncated = clean_extracted_text("\n".join(lines), max_chars=self.max_chars)
        return ExtractedText(
            text=text,
            method="digital" if text else "none",
            truncated=truncated,
            error=None if text else "No readable text found in the spreadsheet.",
        )


_extractor: DocumentExtractor | None = None


def get_extractor() -> DocumentExtractor:
    global _extractor
    if _extractor is None:
        _extractor = DocumentExtractor()
    return _extractor


def reset_extractor() -> None:
    """Drop the cached extractor — used by tests that change settings."""
    global _extractor
    _extractor = None
