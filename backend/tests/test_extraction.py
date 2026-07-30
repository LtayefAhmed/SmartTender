"""Document text extraction and its effect on scoring.

The point of this module is the last class: extraction is not a nice-to-have,
it is what makes scoring work at all on a PDF tender. The tests build real PDFs
and DOCX files rather than mocking the libraries, because the failure modes
that matter — a hyphen split across lines, a requirements table, an empty text
layer — only exist in real files.
"""

from __future__ import annotations

import io
import zipfile

import pytest

from app.services.extraction import (
    DocumentExtractor,
    clean_extracted_text,
    get_extractor,
    reset_extractor,
)


# ---------------------------------------------------------------------------
# Helpers that build genuine documents
# ---------------------------------------------------------------------------
def _pdf_with_text(lines: list[str]) -> bytes:
    """A minimal but genuinely valid PDF carrying a text layer."""
    from pypdf import PdfWriter

    try:
        from reportlab.lib.pagesizes import A4  # noqa: F401
    except ImportError:
        # Build the PDF by hand rather than depending on reportlab.
        content = "BT /F1 12 Tf 50 780 Td 14 TL\n"
        for line in lines:
            escaped = line.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
            content += f"({escaped}) Tj T*\n"
        content += "ET"
        stream = content.encode("latin-1", "replace")

        objects = [
            b"<< /Type /Catalog /Pages 2 0 R >>",
            b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] "
            b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
            b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
            b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        ]

        out = bytearray(b"%PDF-1.4\n")
        offsets = [0]
        for index, body in enumerate(objects, start=1):
            offsets.append(len(out))
            out += f"{index} 0 obj\n".encode() + body + b"\nendobj\n"
        xref_at = len(out)
        out += f"xref\n0 {len(objects) + 1}\n".encode()
        out += b"0000000000 65535 f \n"
        for offset in offsets[1:]:
            out += f"{offset:010d} 00000 n \n".encode()
        out += (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_at}\n%%EOF\n"
        ).encode()
        return bytes(out)

    writer = PdfWriter()  # pragma: no cover - reportlab path unused
    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def _pdf_without_text_layer() -> bytes:
    """A structurally valid PDF whose page carries no text — i.e. a scan."""
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Contents 4 0 R >>",
        b"<< /Length 0 >>\nstream\n\nendstream",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{index} 0 obj\n".encode() + body + b"\nendobj\n"
    xref_at = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode() + b"0000000000 65535 f \n"
    for offset in offsets[1:]:
        out += f"{offset:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_at}\n%%EOF\n"
    ).encode()
    return bytes(out)


def _docx_with(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    import docx

    document = docx.Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if table_rows:
        table = document.add_table(rows=0, cols=len(table_rows[0]))
        for row in table_rows:
            cells = table.add_row().cells
            for cell, value in zip(cells, row, strict=False):
                cell.text = value
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.fixture()
def extractor() -> DocumentExtractor:
    reset_extractor()
    return DocumentExtractor()


# ---------------------------------------------------------------------------
class TestPdfExtraction:
    def test_a_digital_pdf_yields_its_text(self, extractor):
        pdf = _pdf_with_text(
            [
                "Appel d'offres pour le developpement d'une plateforme",
                "de gestion documentaire et la maintenance applicative.",
            ]
        )
        result = extractor.extract(pdf, content_type="application/pdf", filename="cdc.pdf")

        assert result.ok
        assert result.method == "digital"
        assert "developpement" in result.text.lower()
        assert "maintenance applicative" in result.text.lower()

    def test_page_count_is_reported(self, extractor):
        result = extractor.extract(_pdf_with_text(["une ligne"]), filename="x.pdf")
        assert result.pages_total == 1

    def test_a_scanned_pdf_is_routed_to_ocr(self, extractor):
        """A PDF with no text layer must be *detected*, whether or not the
        Tesseract binary happens to be installed on this machine."""
        result = extractor.extract(_pdf_without_text_layer(), filename="scan.pdf")

        if result.method == "ocr":
            pass  # Tesseract present and it read something
        else:
            # Tesseract absent: the tender degrades to no text, with a warning
            # naming the cause — never a crash.
            assert result.method == "none"
            assert not result.ok
            assert result.error

    def test_a_corrupt_pdf_degrades_rather_than_raises(self, extractor):
        result = extractor.extract(b"%PDF-1.4\nthis is not really a pdf", filename="x.pdf")
        assert not result.ok
        assert result.error

    def test_an_oversized_document_is_skipped(self, extractor):
        extractor.max_document_bytes = 100
        result = extractor.extract(_pdf_with_text(["texte"]), filename="big.pdf")
        assert not result.ok
        assert "limit" in result.error


class TestDocxExtraction:
    def test_paragraphs_are_extracted(self, extractor):
        docx_bytes = _docx_with(
            ["Objet du marche : developpement applicatif", "Duree : 36 mois"]
        )
        result = extractor.extract(docx_bytes, filename="cdc.docx")

        assert result.ok
        assert "developpement applicatif" in result.text

    def test_tables_are_extracted(self, extractor):
        """Tender requirements live in tables far more often than in prose, so
        skipping tables would lose the most valuable content in the file."""
        docx_bytes = _docx_with(
            ["Exigences techniques"],
            table_rows=[
                ["Competence", "Niveau"],
                ["Kubernetes", "Expert"],
                ["PostgreSQL", "Confirme"],
            ],
        )
        result = extractor.extract(docx_bytes, filename="exigences.docx")

        assert "Kubernetes" in result.text
        assert "Expert" in result.text

    def test_a_corrupt_docx_degrades(self, extractor):
        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, "w") as archive:
            archive.writestr("word/document.xml", "not valid xml at all <<<")
        result = extractor.extract(buffer.getvalue(), filename="broken.docx")
        assert not result.ok


class TestHtmlExtraction:
    def test_text_is_extracted_and_chrome_removed(self, extractor):
        html = (
            b"<html><head><style>.x{color:red}</style></head><body>"
            b"<nav>Accueil Contact</nav>"
            b"<h1>Avis d'appel d'offres</h1>"
            b"<p>Maintenance applicative du systeme RH.</p>"
            b"<script>track()</script><footer>Mentions legales</footer>"
            b"</body></html>"
        )
        result = extractor.extract(html, filename="avis.html")

        assert "Maintenance applicative" in result.text
        assert "track()" not in result.text
        assert "Mentions legales" not in result.text


class TestTextCleaning:
    def test_hyphenation_across_lines_is_rejoined(self):
        """Without this, "développe-\\nment" never matches the keyword
        "développement" — and keywords carry 15% of the score."""
        text, _ = clean_extracted_text("le developpe-\nment applicatif")
        assert "developpement applicatif" in text

    def test_pdf_ligatures_are_normalised(self):
        text, _ = clean_extracted_text("eﬃcacité et ﬁnancement")
        assert "efficacite" in text.replace("é", "e")
        assert "financement" in text

    def test_typographic_quotes_are_normalised(self):
        text, _ = clean_extracted_text("l’appel d’offres")
        assert "l'appel d'offres" in text

    def test_whitespace_is_collapsed(self):
        text, _ = clean_extracted_text("a    b\n\n\n\nc")
        assert text == "a b\n\nc"

    def test_truncation_is_reported_and_respects_word_boundaries(self):
        text, truncated = clean_extracted_text("mot " * 500, max_chars=100)
        assert truncated is True
        assert len(text) <= 100
        assert not text.endswith("mo")

    def test_empty_input_is_safe(self):
        assert clean_extracted_text("") == ("", False)


class TestExtractionChangesScoring:
    """The reason this feature exists."""

    def test_a_pdf_tender_is_scored_on_its_content_not_its_title(self, make_tender):
        from app.services.scoring import ScoringEngine

        engine = ScoringEngine()

        # A deliberately uninformative title, exactly as portals publish them.
        bare = make_tender(
            title="AO 42/2026",
            description=None,
            sector=None,
            category=None,
            cpv_codes=[],
        )
        without_text = engine.score(bare)

        with_text = engine.score(
            bare.model_copy(
                update={
                    "full_text": (
                        "Le present marche porte sur le developpement d'applications "
                        "web, l'integration au systeme d'information existant, la "
                        "migration des donnees et une tierce maintenance applicative. "
                        "Un consultant expert Java et deux developpeurs confirmes sont "
                        "attendus."
                    )
                }
            )
        )

        assert with_text.score > without_text.score
        assert with_text.breakdown["field_of_work"]["value"] > (
            without_text.breakdown["field_of_work"]["value"] or 0
        )
        assert with_text.breakdown["keywords"]["details"]["matched"]

    def test_extracted_text_does_not_pollute_duplicate_detection(self, make_tender):
        """full_text must stay out of the dedup comparison key, or two unrelated
        tenders sharing boilerplate annexes would look identical."""
        tender = make_tender()
        with_text = tender.model_copy(update={"full_text": "x" * 50_000})

        assert with_text.comparison_text() == tender.comparison_text()
        assert "x" * 100 not in with_text.comparison_text()

    def test_a_blocking_keyword_inside_the_document_still_vetoes(self, make_tender):
        from app.services.scoring import ScoringEngine

        result = ScoringEngine().score(
            make_tender(
                title="AO 51/2026",
                description=None,
                full_text="Le marche porte sur des travaux de genie civil et de gros oeuvre.",
            )
        )
        assert result.veto_reason is not None


class TestExtractorSingleton:
    def test_the_extractor_is_cached_and_resettable(self):
        reset_extractor()
        first = get_extractor()
        assert get_extractor() is first
        reset_extractor()
        assert get_extractor() is not first
